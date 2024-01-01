import os
import re
from docx import Document
import fitz  # PyMuPDF
from striprtf.striprtf import rtf_to_text
import logging
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.pagesizes import letter

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

def is_valid_file_path(path):
    """
    Validates the given file path.

    Args:
        path (str): The file path to validate.

    Returns:
        bool: True if the path is valid and the file exists, False otherwise.
    """
    pattern = r'^[a-zA-Z0-9_\\-\\/:. ]+$'
    if not re.match(pattern, path):
        logging.warning(f"Invalid file path format: {path}")
        return False
    if not os.path.isfile(path):
        logging.warning(f"File does not exist: {path}")
        return False
    return True

def load_document(file_path):
    """
    Loads a document from the given file path.

    Args:
        file_path (str): The path of the file to load.

    Returns:
        str: The text content of the document.

    Raises:
        ValueError: If the file path is invalid or the file extension is unsupported.
    """
    if not is_valid_file_path(file_path):
        raise ValueError(f"Invalid file path: {file_path}")

    file_extension = os.path.splitext(file_path)[1].lower()
    text = ""

    try:
        if file_extension == ".pdf":
            with fitz.open(file_path) as doc:
                text = "\n".join(page.get_text() for page in doc)
        elif file_extension == ".docx":
            doc = Document(file_path)
            text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
        elif file_extension == ".rtf":
            with open(file_path, 'r', encoding='utf-8') as file:
                rtf_text = file.read()
            text = rtf_to_text(rtf_text)
        elif file_extension == ".txt":
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
        else:
            raise ValueError(f"Unsupported file extension: {file_extension}")
    except Exception as e:
        logging.error(f"Error loading document from {file_path}: {e}")
        raise

    logging.info(f"Document loaded successfully from {file_path}")
    return text

def save_summary(summary, file_path):
    """
    Saves the summary to a file.

    Args:
        summary (str): The summary text to save.
        file_path (str): The path where the summary should be saved.

    Raises:
        RuntimeError: If an error occurs during file saving.
    """
    try:
        if file_path.endswith('.txt'):
            with open(file_path, 'w', encoding='utf-8') as file:
                file.write(summary)
        elif file_path.endswith('.docx'):
            doc = Document()
            doc.add_paragraph(summary)
            doc.save(file_path)
        elif file_path.endswith('.pdf'):
            pdf = SimpleDocTemplate(file_path, pagesize=letter)
            styles = getSampleStyleSheet()
            summary_paragraph = Paragraph(summary, styles['Normal'])
            pdf.build([summary_paragraph])
        else:
            raise ValueError("Unsupported file format selected.")
    except Exception as e:
        logging.error(f"An error occurred while saving the file to {file_path}: {e}")
        raise RuntimeError(f"An error occurred while saving the file: {e}")

    logging.info(f"Summary saved successfully to {file_path}")
